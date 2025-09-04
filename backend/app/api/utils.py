from fastapi import APIRouter, HTTPException, status, Depends, Query
from fastapi.responses import FileResponse
import re
import os
import tempfile
from typing import Literal, Optional, List
from pathlib import Path
import json
from app.models import APIResponse, ReformatRequest, ReformatResponse
from app.auth import get_current_user
from app.services.file_service import file_service
from datetime import datetime
from app.services.reformat_service import reformat_service
import logging

logger = logging.getLogger(__name__)

router = APIRouter()

@router.post("/reformat", response_model=APIResponse)
async def reformat_text(
    request: dict,
    current_user: dict = Depends(get_current_user)
):
    """自动排版文本（统一到 reformat_service 实现）"""
    try:
        text = request.get("text", "")
        settings = request.get("settings", {})
        if not text:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="文本内容不能为空"
            )

        formatted_text, diff_info = reformat_service.reformat_text(text, settings)

        return APIResponse(
            ok=True,
            data={
                "formatted_text": formatted_text,
                "diff_info": diff_info
            }
        )
    except HTTPException as e:
        return APIResponse(
            ok=False,
            error={"code": str(e.status_code), "msg": e.detail}
        )
    except Exception as e:
        return APIResponse(
            ok=False,
            error={"code": "500", "msg": f"排版失败: {str(e)}"}
        )

def _format_chinese_text(text: str) -> str:
    """格式化中文文本"""
    if not text:
        return text
    
    # 1. 统一换行符为 \n
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    
    # 2. 压缩连续空行为单个空行（保留preserve标记的除外）
    lines = text.split('\n')
    formatted_lines = []
    prev_empty = False
    
    for line in lines:
        line = line.strip()
        
        # 检查是否有preserve标记
        if '<!--preserve-->' in line:
            formatted_lines.append(line.replace('<!--preserve-->', ''))
            prev_empty = False
            continue
        
        if not line:  # 空行
            if not prev_empty:
                formatted_lines.append('')
                prev_empty = True
        else:
            formatted_lines.append(line)
            prev_empty = False
    
    # 3. 重新组合文本
    text = '\n'.join(formatted_lines)
    
    # 4. 替换省略号
    text = re.sub(r'\.{3,}', '……', text)  # 3个或以上点替换为省略号
    text = re.sub(r'…{2,}', '……', text)   # 多个省略号替换为两个
    
    # 5. 段落首行缩进（两个中文空格）
    lines = text.split('\n')
    formatted_lines = []
    
    for line in lines:
        if line.strip():  # 非空行
            # 移除现有的缩进
            line = line.lstrip()
            # 添加标准缩进（两个中文空格）
            if line and not line.startswith(('　', ' ', '\t')):
                line = '　　' + line
            formatted_lines.append(line)
        else:
            formatted_lines.append('')
    
    # 6. 中英文之间智能空格（可选功能）
    text = '\n'.join(formatted_lines)
    # text = _add_spaces_between_chinese_english(text)
    
    # 7. 清理多余的空白字符
    text = re.sub(r'[ \t]+', ' ', text)  # 多个空格替换为单个空格
    text = re.sub(r' +\n', '\n', text)   # 行末空格
    text = re.sub(r'\n +', '\n', text)   # 行首空格（除了缩进）
    
    return text

def _add_spaces_between_chinese_english(text: str) -> str:
    """在中英文之间添加空格"""
    # 中文字符后跟英文字符
    text = re.sub(r'([\u4e00-\u9fff])([a-zA-Z0-9])', r'\1 \2', text)
    # 英文字符后跟中文字符
    text = re.sub(r'([a-zA-Z0-9])([\u4e00-\u9fff])', r'\1 \2', text)
    return text

@router.post("/export/{novel_id}")
async def export_novel(
    novel_id: str,
    format: Literal["txt", "docx"] = "txt",
    current_user: dict = Depends(get_current_user)
):
    """导出小说为TXT或Word格式"""
    try:
        # 检查小说是否存在
        config = file_service.get_novel_config(novel_id)
        if not config:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="小说不存在"
            )
        
        # 获取章节列表
        chapters_data = file_service.list_chapters(novel_id)
        
        if not chapters_data:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="小说没有章节内容可导出"
            )
        
        # 读取所有章节内容
        chapters = []
        for chapter_file in chapters_data:
            filename = chapter_file["filename"]
            content = file_service.read_chapter(novel_id, filename)
            if content:
                # 解析章节信息
                order, title = _parse_chapter_filename(filename)
                chapters.append({
                    "order": order,
                    "title": title,
                    "content": content
                })
        
        # 按order排序
        chapters.sort(key=lambda x: x["order"])
        
        # 生成导出内容
        novel_title = config.get("title", novel_id)
        export_content = _generate_export_content(novel_title, chapters, config)
        
        # 创建临时文件
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{novel_title}_{timestamp}"
        
        if format == "txt":
            temp_file = tempfile.NamedTemporaryFile(
                mode='w', 
                suffix='.txt', 
                delete=False, 
                encoding='utf-8'
            )
            temp_file.write(export_content)
            temp_file.close()
            
            return FileResponse(
                path=temp_file.name,
                filename=f"{filename}.txt",
                media_type="text/plain",
                background=_cleanup_temp_file(temp_file.name)
            )
        
        elif format == "docx":
            try:
                from docx import Document
                from docx.shared import Inches
                
                # 创建Word文档
                doc = Document()
                
                # 添加标题
                title = doc.add_heading(novel_title, 0)
                title.alignment = 1  # 居中对齐
                
                # 添加小说信息
                if config.get("meta", {}).get("description"):
                    doc.add_paragraph(config["meta"]["description"])
                    doc.add_paragraph("")  # 空行
                
                # 添加章节内容
                for chapter in chapters:
                    # 章节标题
                    doc.add_heading(chapter["title"], level=1)
                    
                    # 章节内容
                    paragraphs = chapter["content"].split('\n')
                    for paragraph in paragraphs:
                        if paragraph.strip():
                            doc.add_paragraph(paragraph.strip())
                    
                    # 章节间空行
                    doc.add_paragraph("")
                
                # 保存到临时文件
                temp_file = tempfile.NamedTemporaryFile(
                    suffix='.docx', 
                    delete=False
                )
                doc.save(temp_file.name)
                temp_file.close()
                
                return FileResponse(
                    path=temp_file.name,
                    filename=f"{filename}.docx",
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    background=_cleanup_temp_file(temp_file.name)
                )
                
            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Word导出功能需要安装python-docx库"
                )
        
        else:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="不支持的导出格式"
            )
            
    except HTTPException as e:
        raise e
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"导出失败: {str(e)}"
        )

def _parse_chapter_filename(filename: str):
    """解析章节文件名，提取order和title"""
    name_without_ext = filename.replace('.txt', '')
    parts = name_without_ext.split('_', 1)
    if len(parts) >= 2 and parts[0].isdigit():
        return int(parts[0]), parts[1]
    return 1, name_without_ext

def _generate_export_content(novel_title: str, chapters: list, config: dict) -> str:
    """生成导出内容"""
    lines = []
    
    # 小说标题
    lines.append(novel_title)
    lines.append("=" * len(novel_title))
    lines.append("")
    
    # 小说信息
    meta = config.get("meta", {})
    if meta.get("author"):
        lines.append(f"作者：{meta['author']}")
    if meta.get("genre"):
        lines.append(f"类型：{meta['genre']}")
    if meta.get("description"):
        lines.append(f"简介：{meta['description']}")
    
    if meta:
        lines.append("")
        lines.append("-" * 50)
        lines.append("")
    
    # 章节内容
    for chapter in chapters:
        lines.append(f"第{chapter['order']}章 {chapter['title']}")
        lines.append("")
        
        # 章节内容，每段缩进
        paragraphs = chapter['content'].split('\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                # 确保段落首行缩进
                paragraph = paragraph.strip()
                if not paragraph.startswith('　'):
                    paragraph = '　　' + paragraph
                lines.append(paragraph)
            else:
                lines.append("")
        
        lines.append("")
        lines.append("")
    
    return '\n'.join(lines)

def _cleanup_temp_file(file_path: str):
    """清理临时文件的后台任务"""
    import asyncio
    
    async def cleanup():
        await asyncio.sleep(1)  # 等待文件下载完成
        try:
            if os.path.exists(file_path):
                os.unlink(file_path)
        except:
            pass
    
    return cleanup()

@router.get("/health", response_model=APIResponse)
async def health_check():
    """健康检查"""
    return APIResponse(
        ok=True,
        data={
            "status": "healthy",
            "message": "工具服务运行正常",
            "available_tools": ["reformat", "export", "prompts"]
        }
    ) 

# ---- AI 提示词读写 ----

_DEFAULT_PROMPTS = {
    "expand": {
        "sentence": "请自然地扩写这句话，使其更生动但不改变原意。",
        "paragraph": "请根据以下段落进行段落级扩写，补充细节、动作与情绪，保持叙事连贯。"
    },
    "polish": {
        "sentence": "在不改变含义的前提下，润色这句话，使其更流畅有表现力。",
        "paragraph": "请润色下文段落，优化节奏与措辞，必要时调整句式提升可读性。"
    },
    "summary": {
        "chapter": "请生成本章节的概要（3-5句），突出关键冲突、转折与推进。"
    },
    "world": {
        "random": "随机生成一个世界观元素（设定/地点/历史事件），要求简洁、有画面感，可落地。",
        "targeted": "围绕下文设定进行细化（地理/势力/规则/禁忌），保持一致性与可用性。"
    },
    "character": {
        "single": "请基于提供的文本，整理该角色卡片（姓名/身份/性格/动机/欲望/冲突/人际/关键道具/成长线）。",
        "all": "请通读整篇大纲，整理主要角色卡片列表（每个含：姓名/身份/性格/动机/关系/弧光）。"
    },
    "outline": {
        "story_background": {
            "random": "生成故事背景设定，包含世界观概述、核心设定、地理环境、社会结构等。",
            "targeted": "基于已有内容生成故事背景设定，保持一致性。"
        },
        "power_system": {
            "random": "生成力量体系，包含修炼体系、技能系统、限制规则、成长路径等。",
            "targeted": "基于已有内容生成力量体系，保持一致性。"
        },
        "history_background": {
            "random": "生成历史背景，包含重要历史事件、历史人物、历史影响、历史脉络等。",
            "targeted": "基于已有内容生成历史背景，保持一致性。"
        },
        "story_timeline": {
            "random": "生成故事时间线，包含整体时间跨度、关键时间节点、时间管理、时间节奏等。",
            "targeted": "基于已有内容生成故事时间线，保持一致性。"
        },
        "story_location": {
            "random": "生成故事发生地点，包含主要场景、地理特征、人文环境、场景氛围等。",
            "targeted": "基于已有内容生成故事发生地点，保持一致性。"
        },
        "main_plot": {
            "random": "生成故事主线，包含核心冲突、情节发展、主题思想、人物弧光等。",
            "targeted": "基于已有内容生成故事主线，保持一致性。"
        },
        "volume_summary": {
            "random": "生成分卷内容简介，包含分卷规划、每卷内容、分卷关联等。",
            "targeted": "基于已有内容生成分卷内容简介，保持一致性。"
        }
    },
    "inspiration": "你是资深畅销网文作家。基于类型、篇幅、标签与偏好，产出300-600字主体故事灵感，张力足、节奏明快、可扩写为大纲。"
}

def _deep_merge(a: dict, b: dict) -> dict:
    result = dict(a)
    for k, v in (b or {}).items():
        if isinstance(v, dict) and isinstance(result.get(k), dict):
            result[k] = _deep_merge(result[k], v)
        else:
            result[k] = v
    return result

def _normalize_prompts(data: dict) -> dict:
    """兼容旧版扁平结构到新版分层结构。"""
    if not data or any(isinstance(data.get(k), dict) for k in ("expand","polish","summary","world","character")):
        # 认为已是新版或空
        return _deep_merge(_DEFAULT_PROMPTS, data or {})
    # 旧版到新版映射
    mapped = {
        "expand": {
            "sentence": data.get("expand") or _DEFAULT_PROMPTS["expand"]["sentence"],
            "paragraph": data.get("expand") or _DEFAULT_PROMPTS["expand"]["paragraph"],
        },
        "polish": {
            "sentence": data.get("polish") or _DEFAULT_PROMPTS["polish"]["sentence"],
            "paragraph": data.get("polish") or _DEFAULT_PROMPTS["polish"]["paragraph"],
        },
        "summary": {
            "chapter": data.get("chapter_summary") or data.get("summarize") or _DEFAULT_PROMPTS["summary"]["chapter"]
        },
        "world": {
            "random": data.get("world_random") or _DEFAULT_PROMPTS["world"]["random"],
            "targeted": data.get("world_targeted") or _DEFAULT_PROMPTS["world"]["targeted"],
        },
        "character": {
            "single": data.get("character_design") or _DEFAULT_PROMPTS["character"]["single"],
            "all": data.get("character_design") or _DEFAULT_PROMPTS["character"]["all"],
        },
        "inspiration": data.get("inspiration") or _DEFAULT_PROMPTS["inspiration"],
    }
    return mapped

def _normalize_prompts_partial(data: dict) -> dict:
    """仅规范化请求体中提供的字段，不注入默认值。
    - 支持新版分层字段的透传
    - 支持旧版扁平字段到分层字段的映射（仅映射出现的键）
    """
    if not isinstance(data, dict):
        return {}
    out: dict = {}
    # 直接透传分层结构里出现的子键
    if isinstance(data.get("expand"), dict):
        out.setdefault("expand", {})
        if isinstance(data["expand"].get("sentence"), str):
            out["expand"]["sentence"] = data["expand"]["sentence"]
        if isinstance(data["expand"].get("paragraph"), str):
            out["expand"]["paragraph"] = data["expand"]["paragraph"]
        if not out["expand"]:
            out.pop("expand")
    elif isinstance(data.get("expand"), str):
        out["expand"] = {"sentence": data["expand"], "paragraph": data["expand"]}

    if isinstance(data.get("polish"), dict):
        out.setdefault("polish", {})
        if isinstance(data["polish"].get("sentence"), str):
            out["polish"]["sentence"] = data["polish"]["sentence"]
        if isinstance(data["polish"].get("paragraph"), str):
            out["polish"]["paragraph"] = data["polish"]["paragraph"]
        if not out["polish"]:
            out.pop("polish")
    elif isinstance(data.get("polish"), str):
        out["polish"] = {"sentence": data["polish"], "paragraph": data["polish"]}

    # summary / chapter_summary 兼容
    if isinstance(data.get("summary"), dict) and isinstance(data["summary"].get("chapter"), str):
        out["summary"] = {"chapter": data["summary"]["chapter"]}
    elif isinstance(data.get("chapter_summary"), str):
        out["summary"] = {"chapter": data.get("chapter_summary")}

    # world
    if isinstance(data.get("world"), dict):
        out.setdefault("world", {})
        if isinstance(data["world"].get("random"), str):
            out["world"]["random"] = data["world"]["random"]
        if isinstance(data["world"].get("targeted"), str):
            out["world"]["targeted"] = data["world"]["targeted"]
        if not out["world"]:
            out.pop("world")
    else:
        if isinstance(data.get("world_random"), str) or isinstance(data.get("world_targeted"), str):
            out["world"] = {}
            if isinstance(data.get("world_random"), str):
                out["world"]["random"] = data["world_random"]
            if isinstance(data.get("world_targeted"), str):
                out["world"]["targeted"] = data["world_targeted"]

    # character
    if isinstance(data.get("character"), dict):
        out.setdefault("character", {})
        if isinstance(data["character"].get("single"), str):
            out["character"]["single"] = data["character"]["single"]
        if isinstance(data["character"].get("all"), str):
            out["character"]["all"] = data["character"]["all"]
        if not out["character"]:
            out.pop("character")
    else:
        if isinstance(data.get("character_design"), str):
            out["character"] = {"single": data["character_design"], "all": data["character_design"]}

    # inspiration
    if isinstance(data.get("inspiration"), str):
        out["inspiration"] = data["inspiration"]

    # styles
    if isinstance(data.get("styles"), dict):
        out["styles"] = {}
        if isinstance(data["styles"].get("current"), str):
            out["styles"]["current"] = data["styles"]["current"]
        if isinstance(data["styles"].get("history"), list):
            out["styles"]["history"] = [s for s in data["styles"]["history"] if isinstance(s, str)]
        if not out["styles"]:
            out.pop("styles", None)

    return out

def _prompts_file(scope: str, novel_id: str | None) -> Path:
    base = Path(file_service.novel_repo_path)
    if scope == "novel" and novel_id:
        return base / novel_id / "ai_prompts.json"
    return base / "ai_prompts.json"

def _character_prompts_file() -> Path:
    base = Path(file_service.novel_repo_path)
    return base / "ai_character_prompts.json"

def _character_prompts_dir() -> Path:
    base = Path(file_service.novel_repo_path)
    return base / "ai_prompts" / "character"

@router.get("/prompts")
async def get_prompts(
    scope: str = Query(..., description="提示词范围: global/novel"),
    novel_id: Optional[str] = Query(None, description="小说ID（novel范围时必填）"),
    current_user: dict = Depends(get_current_user)
):
    """获取AI提示词配置"""
    try:
        if scope == "novel" and not novel_id:
            return APIResponse(ok=False, error={"code": "400", "msg": "小说范围需要提供novel_id"})
        
        if scope == "global":
            # 读取全局提示词 - 新的分散文件结构（使用绝对路径）
            base_dir = Path(file_service.novel_repo_path)
            prompts_dir = base_dir / "ai_prompts"
            character_prompts_file = base_dir / "ai_character_prompts.json"
            
            # 调试信息
            logger.info(f"Reading prompts from base_dir: {base_dir}")
            logger.info(f"Prompts directory: {prompts_dir}")
            logger.info(f"Character prompts file: {character_prompts_file}")
        else:
            # 读取小说级提示词（使用绝对路径）
            base_dir = Path(file_service.novel_repo_path) / str(novel_id)
            prompts_dir = base_dir / "ai_prompts"
            character_prompts_file = base_dir / "ai_character_prompts.json"
        
        prompts = {}
        character_prompts = {}
        outline_writing_prompts = {}
        
        # 读取分散的提示词文件（一级：expand/polish/summary/world/inspiration/styles 等）
        if Path(prompts_dir).exists():
            logger.info(f"Prompts directory exists: {prompts_dir}")
            try:
                # 遍历ai_prompts目录下的所有json文件
                for fp in Path(prompts_dir).iterdir():
                    if fp.is_file() and fp.suffix == '.json':
                        prompt_type = fp.stem  # 去掉.json后缀
                        logger.info(f"Found prompt file: {fp} -> {prompt_type}")
                        try:
                            with open(fp, 'r', encoding='utf-8') as f:
                                prompt_data = json.load(f)
                                prompts[prompt_type] = prompt_data
                                logger.info(f"Loaded {prompt_type}: {list(prompt_data.keys()) if isinstance(prompt_data, dict) else 'not dict'}")
                        except json.JSONDecodeError as e:
                            logger.error(f"JSON解析错误 in {fp}: {e}")
                            # 如果JSON解析失败，使用默认值
                            if prompt_type == "expand":
                                prompts[prompt_type] = _DEFAULT_PROMPTS.get("expand", {})
                            elif prompt_type == "polish":
                                prompts[prompt_type] = _DEFAULT_PROMPTS.get("polish", {})
                            elif prompt_type == "summary":
                                prompts[prompt_type] = _DEFAULT_PROMPTS.get("summary", {})
                            elif prompt_type == "world":
                                prompts[prompt_type] = _DEFAULT_PROMPTS.get("world", {})
                            elif prompt_type == "inspiration":
                                prompts[prompt_type] = _DEFAULT_PROMPTS.get("inspiration", "")
                            elif prompt_type == "styles":
                                prompts[prompt_type] = {"current": "默认"}
                        except Exception as e:
                            logger.error(f"读取提示词文件失败 {fp}: {e}")
                            continue
            except Exception as e:
                logger.warning(f"读取提示词目录失败: {e}")
        else:
            logger.warning(f"Prompts directory does not exist: {prompts_dir}")
        
        # 读取大纲编写提示词（outline/*.json）
        try:
            outline_dir = Path(prompts_dir) / "outline"
            if outline_dir.exists():
                for fp in outline_dir.iterdir():
                    if fp.is_file() and fp.suffix == '.json':
                        key = fp.stem
                        try:
                            with open(fp, 'r', encoding='utf-8') as f:
                                data = json.load(f) or {}
                        except Exception:
                            data = {}
                        # 补全默认结构
                        defaults = _DEFAULT_PROMPTS.get("outline", {}).get(key, {}) if key in (_DEFAULT_PROMPTS.get("outline") or {}) else {}
                        merged = {
                            "random": data.get("random") or defaults.get("random", ""),
                            "targeted": data.get("targeted") or defaults.get("targeted", "")
                        }
                        outline_writing_prompts[key] = merged
            # 如果目录不存在但需要返回默认结构
            if not outline_writing_prompts and _DEFAULT_PROMPTS.get("outline"):
                for k, v in _DEFAULT_PROMPTS["outline"].items():
                    outline_writing_prompts[k] = {
                        "random": v.get("random", ""),
                        "targeted": v.get("targeted", "")
                    }
        except Exception as e:
            logger.warning(f"读取大纲编写提示词失败: {e}")

        # 读取角色提示词 - 优先尝试新的分散文件结构
        character_prompts = {}
        
        # 优先尝试新的分散文件结构：novel_repo/ai_prompts/character/角色名.json
        character_dir = Path(prompts_dir) / "character"
        if character_dir.exists():
            logger.info(f"Character directory exists: {character_dir}")
            try:
                # 遍历character目录下的所有json文件
                for fp in character_dir.iterdir():
                    if fp.is_file() and fp.suffix == '.json':
                        role_name = fp.stem
                        logger.info(f"Found character file: {fp} -> {role_name}")
                        try:
                            with open(fp, 'r', encoding='utf-8') as f:
                                role_data = json.load(f)
                                character_prompts[role_name] = role_data
                                logger.info(f"Loaded character {role_name}: {list(role_data.keys()) if isinstance(role_data, dict) else 'not dict'}")
                        except json.JSONDecodeError as e:
                            logger.error(f"JSON解析错误 in character file {fp}: {e}")
                            # 如果JSON解析失败，使用默认值
                            character_prompts[role_name] = {
                                "prompt": f"设计{role_name}角色",
                                "style": "简洁有力",
                                "schema_hint": "输出JSON友好可解析"
                            }
                        except Exception as e:
                            logger.error(f"读取角色提示词文件失败 {fp}: {e}")
                            continue
            except Exception as e:
                logger.warning(f"读取分散角色提示词目录失败: {e}")
        else:
            logger.warning(f"Character directory does not exist: {character_dir}")
        
        # 如果没有读取到分散的角色提示词，回退到旧的文件结构
        if not character_prompts and Path(character_prompts_file).exists():
            try:
                with open(character_prompts_file, 'r', encoding='utf-8') as f:
                    character_prompts = json.load(f)
            except Exception as e:
                logger.warning(f"读取角色提示词文件失败: {e}")
        
        # 确保所有必要的提示词类型都存在，如果不存在则使用默认值
        if "expand" not in prompts:
            prompts["expand"] = _DEFAULT_PROMPTS["expand"]
        if "polish" not in prompts:
            prompts["polish"] = _DEFAULT_PROMPTS["polish"]
        if "summary" not in prompts:
            prompts["summary"] = _DEFAULT_PROMPTS["summary"]
        if "world" not in prompts:
            prompts["world"] = _DEFAULT_PROMPTS["world"]
        if "inspiration" not in prompts:
            prompts["inspiration"] = _DEFAULT_PROMPTS["inspiration"]
        if "styles" not in prompts:
            prompts["styles"] = {"current": "默认"}
        
        # 合并提示词
        all_prompts = {
            **prompts,
            "character": character_prompts,
            "outline_writing": outline_writing_prompts
        }
        
        logger.info(f"Final prompts structure: {list(all_prompts.keys())}")
        logger.info(f"Character prompts keys: {list(character_prompts.keys())}")
        
        return APIResponse(ok=True, data={"prompts": all_prompts})
        
    except Exception as e:
        logger.error(f"获取提示词失败: {e}")
        return APIResponse(ok=False, error={"code": "500", "msg": f"获取提示词失败: {str(e)}"})

@router.post("/prompts")
async def update_prompts(
    request: dict,
    scope: str = Query(..., description="提示词范围: global/novel"),
    novel_id: Optional[str] = Query(None, description="小说ID（novel范围时必填）"),
    current_user: dict = Depends(get_current_user)
):
    """更新AI提示词配置"""
    try:
        if scope == "novel" and not novel_id:
            return APIResponse(ok=False, error={"code": "400", "msg": "小说范围需要提供novel_id"})
        
        if scope == "global":
            base_dir = Path(file_service.novel_repo_path)
            prompts_dir = base_dir / "ai_prompts"
            character_prompts_file = base_dir / "ai_character_prompts.json"
        else:
            # 确保小说目录存在
            base_dir = Path(file_service.novel_repo_path) / str(novel_id)
            base_dir.mkdir(parents=True, exist_ok=True)
            
            prompts_dir = base_dir / "ai_prompts"
            character_prompts_file = base_dir / "ai_character_prompts.json"
        
        # 获取提示词数据 - 直接使用整个请求体
        prompts_data = request
        logger.info(f"Received prompts data: {list(prompts_data.keys())}")

        # 1) 专门处理大纲编写（outline_writing）- 支持两种格式：
        #    A. { category: 'outline_writing', key: 'story_background', type: 'random'|'targeted', value: '...' }
        #    B. { outline_writing: { story_background: { random: '...', targeted: '...' }, ... } }
        try:
            outline_dir = Path(prompts_dir) / "outline"
            if prompts_data.get("category") == "outline_writing":
                key = prompts_data.get("key")
                otype = prompts_data.get("type")  # random | targeted
                value = prompts_data.get("value", "")
                if key and otype in ("random", "targeted"):
                    outline_dir.mkdir(parents=True, exist_ok=True)
                    target_file = outline_dir / f"{key}.json"
                    existing = {}
                    if target_file.exists():
                        try:
                            with open(target_file, 'r', encoding='utf-8') as f:
                                existing = json.load(f)
                        except Exception:
                            existing = {}
                    existing[otype] = value
                    with open(target_file, 'w', encoding='utf-8') as f:
                        json.dump(existing, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"保存大纲编写提示词失败(category形式): {e}")

        # 另一种整体结构写入
        try:
            if isinstance(prompts_data.get("outline_writing"), dict):
                outline_dir.mkdir(parents=True, exist_ok=True)
                for k, v in prompts_data["outline_writing"].items():
                    if not isinstance(v, dict):
                        continue
                    target_file = outline_dir / f"{k}.json"
                    existing = {}
                    if target_file.exists():
                        try:
                            with open(target_file, 'r', encoding='utf-8') as f:
                                existing = json.load(f)
                        except Exception:
                            existing = {}
                    merged = {**existing}
                    if isinstance(v.get("random"), str):
                        merged["random"] = v["random"]
                    if isinstance(v.get("targeted"), str):
                        merged["targeted"] = v["targeted"]
                    with open(target_file, 'w', encoding='utf-8') as f:
                        json.dump(merged, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error(f"保存大纲编写提示词失败(整体结构): {e}")
        
        # 处理角色提示词 - 每个角色单独保存到对应的JSON文件
        if "character" in prompts_data:
            character_prompts = prompts_data["character"]
            logger.info(f"Processing character prompts: {list(character_prompts.keys())}")
            
            # 确保角色提示词目录存在
            character_dir = Path(prompts_dir) / "character"
            character_dir.mkdir(parents=True, exist_ok=True)
            
            # 遍历每个角色，保存到对应的JSON文件
            for role_name, role_data in character_prompts.items():
                if role_data:  # 只保存非空的角色提示词
                    role_file = character_dir / f"{role_name}.json"
                    logger.info(f"Processing role {role_name} -> {role_file}")
                    
                    # 读取现有角色提示词（如果存在）
                    existing_role_data = {}
                    if Path(role_file).exists():
                        try:
                            with open(role_file, 'r', encoding='utf-8') as f:
                                existing_role_data = json.load(f)
                            logger.info(f"Loaded existing data for {role_name}: {list(existing_role_data.keys())}")
                        except Exception as e:
                            logger.warning(f"读取现有角色提示词文件失败 {role_file}: {e}")
                            existing_role_data = {}
                    
                    # 合并角色提示词 - 只更新指定的字段
                    updated_role_data = {**existing_role_data, **role_data}
                    logger.info(f"Updated {role_name} data: {list(updated_role_data.keys())}")
                    
                    # 保存角色提示词到对应的JSON文件
                    try:
                        with open(role_file, 'w', encoding='utf-8') as f:
                            json.dump(updated_role_data, f, ensure_ascii=False, indent=2)
                        logger.info(f"✅ Successfully saved character {role_name} to {role_file}")
                    except Exception as e:
                        logger.error(f"❌ 保存角色提示词文件失败 {role_file}: {e}")
                        raise
                    
                    # 同时保存到旧的文件结构作为备份（可选）
                    try:
                        backup_file = base_dir / "ai_character_prompts.json"
                        with open(backup_file, 'w', encoding='utf-8') as f:
                            json.dump(character_prompts, f, ensure_ascii=False, indent=2)
                        logger.info(f"✅ Successfully saved character prompts backup to {backup_file}")
                    except Exception as e:
                        logger.warning(f"保存角色提示词备份失败: {e}")
        
        # 处理其他类型的提示词（扩写、润色、总结、世界观等）
        other_prompts = {k: v for k, v in prompts_data.items() if k != "character"}
        # 已在上方处理过 outline_writing/category，这里排除
        other_prompts.pop("outline_writing", None)
        other_prompts.pop("category", None)
        if other_prompts:
            logger.info(f"Processing other prompts: {list(other_prompts.keys())}")
            
            # 遍历其他提示词类型
            for prompt_type, prompt_data in other_prompts.items():
                if prompt_data:  # 只保存非空的提示词
                    prompt_file = Path(prompts_dir) / f"{prompt_type}.json"
                    logger.info(f"Processing {prompt_type} -> {prompt_file}")
                    
                    # 读取现有提示词（如果存在）
                    existing_data = {}
                    if Path(prompt_file).exists():
                        try:
                            with open(prompt_file, 'r', encoding='utf-8') as f:
                                existing_data = json.load(f)
                            logger.info(f"Loaded existing data for {prompt_type}: {list(existing_data.keys())}")
                        except Exception as e:
                            logger.warning(f"读取现有提示词文件失败 {prompt_file}: {e}")
                            existing_data = {}
                    
                    # 合并提示词
                    updated_data = {**existing_data, **prompt_data}
                    logger.info(f"Updated {prompt_type} data: {list(updated_data.keys())}")
                    
                    # 保存提示词
                    try:
                        with open(prompt_file, 'w', encoding='utf-8') as f:
                            json.dump(updated_data, f, ensure_ascii=False, indent=2)
                        logger.info(f"✅ Successfully saved {prompt_type} to {prompt_file}")
                    except Exception as e:
                        logger.error(f"❌ 保存提示词文件失败 {prompt_file}: {e}")
                        raise
        
        logger.info(f"Successfully updated prompts for scope: {scope}, novel_id: {novel_id}")
        return APIResponse(ok=True, data={"message": "提示词更新成功"})
        
    except Exception as e:
        logger.error(f"更新提示词失败: {e}")
        return APIResponse(ok=False, error={"code": "500", "msg": f"更新提示词失败: {str(e)}"})

@router.delete("/prompts")
async def delete_prompts(
    scope: str = Query(..., description="提示词范围: global/novel"),
    novel_id: Optional[str] = Query(None, description="小说ID（novel范围时必填）"),
    prompt_keys: List[str] = Query(..., description="要删除的提示词键列表"),
    current_user: dict = Depends(get_current_user)
):
    """删除指定的AI提示词"""
    try:
        if scope == "novel" and not novel_id:
            return APIResponse(ok=False, error={"code": "400", "msg": "小说范围需要提供novel_id"})
        
        if scope == "global":
            prompts_file = "novel_repo/ai_prompts.json"
            character_prompts_file = "novel_repo/ai_character_prompts.json"
        else:
            prompts_file = f"novel_repo/{novel_id}/ai_prompts.json"
            character_prompts_file = f"novel_repo/{novel_id}/ai_character_prompts.json"
        
        deleted_count = 0
        
        # 删除基础提示词
        if os.path.exists(prompts_file):
            try:
                with open(prompts_file, 'r', encoding='utf-8') as f:
                    prompts = json.load(f)
                
                for key in prompt_keys:
                    if key in prompts:
                        del prompts[key]
                        deleted_count += 1
                
                with open(prompts_file, 'w', encoding='utf-8') as f:
                    json.dump(prompts, f, ensure_ascii=False, indent=2)
            except Exception as e:
                logger.warning(f"删除基础提示词失败: {e}")
        
        # 删除角色提示词
        if os.path.exists(character_prompts_file):
            try:
                with open(character_prompts_file, 'r', encoding='utf-8') as f:
                    character_prompts = json.load(f)
                
                for key in prompt_keys:
                    if key in character_prompts:
                        del character_prompts[key]
                        deleted_count += 1
                
                with open(character_prompts_file, 'w', encoding='utf-8') as f:
                    json.dump(character_prompts, f, ensure_ascii=False, indent=2)
            except Exception as e:
                logger.warning(f"删除角色提示词失败: {e}")
        
        return APIResponse(ok=True, data={"message": f"成功删除{deleted_count}个提示词"})
        
    except Exception as e:
        logger.error(f"删除提示词失败: {e}")
        return APIResponse(ok=False, error={"code": "500", "msg": f"删除提示词失败: {str(e)}"})

# ---- 角色提示词（独立文件） ----

_DEFAULT_CHARACTER_PROMPTS = {
    "roles": {
        "女主角": {"prompt": "", "style": ""},
        "男主角": {"prompt": "", "style": ""},
        "女二": {"prompt": "", "style": ""},
        "男二": {"prompt": "", "style": ""},
        "配角": {"prompt": "", "style": ""}
    },
    "schema_hint": "输出JSON友好可解析，字段名固定：appearance/personality/profile(内含:身份职业/家庭关系/早年经历/观念信仰/优点/缺点/成就/社会阶层/习惯嗜好)"
}

@router.get("/character-prompts", response_model=APIResponse)
async def get_character_prompts(current_user: dict = Depends(get_current_user)):
    try:
        # 优先尝试新的分散文件结构
        character_dir = _character_prompts_dir()
        raw = {}
        
        if character_dir.exists():
            try:
                # 遍历character目录下的所有json文件
                for filename in character_dir.iterdir():
                    if filename.is_file() and filename.suffix == '.json':
                        role_name = filename.stem  # 去掉.json后缀
                        with open(filename, 'r', encoding='utf-8') as f:
                            role_data = json.load(f)
                            raw[role_name] = role_data
            except Exception as e:
                logger.warning(f"读取分散角色提示词目录失败: {e}")
        
        # 如果没有读取到分散的角色提示词，回退到旧的文件结构
        if not raw:
            path = _character_prompts_file()
            if path.exists():
                with open(path, 'r', encoding='utf-8') as f:
                    raw = json.load(f)
        
        data = _deep_merge(_DEFAULT_CHARACTER_PROMPTS, raw or {})
        return APIResponse(ok=True, data={"prompts": data})
    except Exception as e:
        return APIResponse(ok=False, error={"code": "500", "msg": f"读取角色提示词失败: {str(e)}"})

@router.post("/character-prompts", response_model=APIResponse)
async def save_character_prompts(payload: dict, current_user: dict = Depends(get_current_user)):
    try:
        updates = payload.get("prompts") or {}
        path = _character_prompts_file()
        character_dir = _character_prompts_dir()
        
        # 确保目录存在
        path.parent.mkdir(parents=True, exist_ok=True)
        character_dir.mkdir(parents=True, exist_ok=True)
        
        existing = {}
        if path.exists():
            try:
                with open(path, 'r', encoding='utf-8') as f:
                    existing = json.load(f)
            except Exception:
                existing = {}
        
        merged = _deep_merge(_DEFAULT_CHARACTER_PROMPTS, _deep_merge(existing, updates))
        
        # 保存到旧的文件结构作为备份
        with open(path, 'w', encoding='utf-8') as f:
            json.dump(merged, f, ensure_ascii=False, indent=2)
        
        # 同时保存到新的分散文件结构
        try:
            for role_name, role_data in merged.get("roles", {}).items():
                if role_data:  # 只保存非空的角色提示词
                    role_file = character_dir / f"{role_name}.json"
                    
                    # 读取现有角色提示词（如果存在）
                    existing_role_data = {}
                    if role_file.exists():
                        try:
                            with open(role_file, 'r', encoding='utf-8') as f:
                                existing_role_data = json.load(f)
                        except Exception:
                            existing_role_data = {}
                    
                    # 合并角色提示词
                    updated_role_data = {**existing_role_data, **role_data}
                    
                    # 保存角色提示词
                    with open(role_file, 'w', encoding='utf-8') as f:
                        json.dump(updated_role_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.warning(f"保存角色提示词到分散文件结构失败: {e}")
        
        return APIResponse(ok=True, data={"prompts": merged})
    except Exception as e:
        return APIResponse(ok=False, error={"code": "500", "msg": f"保存角色提示词失败: {str(e)}"})

@router.post("/reformat", response_model=APIResponse)
async def reformat_text(
    text_data: dict,
    current_user: dict = Depends(get_current_user)
):
    """自动排版文本"""
    try:
        text = text_data.get("text", "")
        settings = text_data.get("settings", {})
        
        if not text:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="文本内容不能为空"
            )
        
        formatted_text, diff_info = reformat_service.reformat_text(text, settings)
        
        return APIResponse(ok=True, data={
            "formatted_text": formatted_text,
            "diff_info": diff_info
        })
    except HTTPException as e:
        return APIResponse(
            ok=False,
            error={"code": str(e.status_code), "msg": e.detail}
        )
    except Exception as e:
        return APIResponse(
            ok=False,
            error={"code": "500", "msg": f"排版失败: {str(e)}"}
        )

@router.get("/export/{novel_id}", response_model=APIResponse)
async def export_novel(
    novel_id: str,
    format: str = "txt",
    current_user: dict = Depends(get_current_user)
):
    """导出小说"""
    try:
        if format not in ["txt", "docx"]:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="不支持的导出格式"
            )
        
        # 获取小说配置
        config = file_service.get_novel_config(novel_id)
        if not config:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="小说不存在"
            )
        
        # 获取所有章节
        chapters = file_service.list_chapters(novel_id)
        if not chapters:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="没有找到章节内容"
            )
        
        # 按顺序排序章节
        chapters.sort(key=lambda x: x.get("order", 0))
        
        # 构建导出内容
        export_content = f"# {config.get('title', 'Untitled')}\n\n"
        
        for chapter in chapters:
            chapter_content = file_service.read_chapter(novel_id, chapter["filename"])
            if chapter_content:
                export_content += f"## {chapter.get('title', 'Untitled')}\n\n"
                export_content += chapter_content + "\n\n"
        
        # 根据格式处理
        if format == "txt":
            # 纯文本格式
            from fastapi.responses import PlainTextResponse
            return PlainTextResponse(
                content=export_content,
                media_type="text/plain; charset=utf-8",
                headers={
                    "Content-Disposition": f"attachment; filename={novel_id}.txt"
                }
            )
        elif format == "docx":
            # Word格式（需要安装python-docx）
            try:
                from docx import Document
                from docx.shared import Inches
                from io import BytesIO
                
                doc = Document()
                doc.add_heading(config.get('title', 'Untitled'), 0)
                
                for chapter in chapters:
                    chapter_content = file_service.read_chapter(novel_id, chapter["filename"])
                    if chapter_content:
                        doc.add_heading(chapter.get('title', 'Untitled'), level=1)
                        doc.add_paragraph(chapter_content)
                
                # 保存到内存
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                from fastapi.responses import StreamingResponse
                return StreamingResponse(
                    iter([buffer.getvalue()]),
                    media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={
                        "Content-Disposition": f"attachment; filename={novel_id}.docx"
                    }
                )
            except ImportError:
                raise HTTPException(
                    status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                    detail="Word导出功能需要安装python-docx库"
                )
        
    except HTTPException as e:
        return APIResponse(
            ok=False,
            error={"code": str(e.status_code), "msg": e.detail}
        )
    except Exception as e:
        return APIResponse(
            ok=False,
            error={"code": "500", "msg": f"导出失败: {str(e)}"}
        )
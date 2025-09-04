import sys
from pathlib import Path

from docx import Document


def docx_to_text(docx_path: Path) -> str:
    document = Document(str(docx_path))
    lines = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            lines.append("\t".join(cells))
    return "\n".join(lines)


def convert(paths: list[Path]) -> None:
    for p in paths:
        txt_path = p.with_suffix(".txt")
        content = docx_to_text(p)
        with txt_path.open("w", encoding="utf-8", newline="\n") as f:
            f.write(content)


def main() -> None:
    base = Path(__file__).resolve().parents[1]
    default_paths = [
        base / "资料/0-AI写作汇总/【短中长篇小说】AI辅助小说创作指令+教程/1.男主人设灵感+完善+整体人设生成/1.男主人设灵感-AI指令_02_07_223902.docx",
        base / "资料/0-AI写作汇总/【短中长篇小说】AI辅助小说创作指令+教程/2.女主人设灵感+完善+整体人设生成/2.女主人设灵感-AI指令_02_07_224602.docx",
        base / "资料/0-AI写作汇总/【短中长篇小说】AI辅助小说创作指令+教程/3.故事亮点灵感(时间空间人设火花立意)/3.故事亮点灵感(时间空间人设火花)-AI指令_03_07_074707.docx",
        base / "资料/0-AI写作汇总/【短中长篇小说】AI辅助小说创作指令+教程/4.世界架构灵感/4.世界架构灵感-AI指令_03_09_205703.docx",
        base / "资料/0-AI写作汇总/【短中长篇小说】AI辅助小说创作指令+教程/5.爽点热梗灵感/5.爽点热梗灵感-AI指令_03_12_234730.docx",
        base / "资料/0-AI写作汇总/【短中长篇小说】AI辅助小说创作指令+教程/7.配角人设生成/7.配角人设生成灵感-AI指令_03_23_190519.docx",
        base / "资料/0-AI写作汇总/【短中长篇小说】AI辅助小说创作指令+教程/8.角色间亮点、群像情节设计/8.角色间亮点、群像情节灵感-AI指令_04_04_162208.docx",
    ]

    if len(sys.argv) > 1:
        paths = [Path(a) for a in sys.argv[1:]]
    else:
        paths = default_paths

    existing = [p for p in paths if p.exists()]
    if existing:
        convert(existing)


if __name__ == "__main__":
    main()


